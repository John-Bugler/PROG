import os
from docx import Document
from docx.shared import Cm, Inches
from pdf2docx import parse
import fitz  # PyMuPDF

# PDF jako text
def create_docx_with_pdf_content(pdf_file, docx_file):  
    # Vytvoření prázdného dokumentu DOCX s požadovanými okraji
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)  # Okraj nahoře 2.5 cm
        section.bottom_margin = Cm(2.5)  # Okraj dole 2.5 cm
        section.left_margin = Cm(2)  # Okraj vlevo 2 cm
        section.right_margin = Cm(2)  # Okraj vpravo 2 cm
    doc.save(docx_file)

    # Získání textu z PDF
    pdf_text = parse(pdf_file)

    # Přidání textu do vytvořeného DOCX souboru
    doc = Document(docx_file)
    for paragraph in pdf_text:
        doc.add_paragraph(paragraph)
    doc.save(docx_file)

# PDF jako obrazek
def insert_pdf_pages_as_images(pdf_file, docx_file):   
    # Inicializace dokumentu DOCX
    doc = Document()

    # Otevření PDF souboru
    pdf_document = fitz.open(pdf_file)

    # Procházení stránek PDF a vložení jako obrázky do dokumentu DOCX
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        img_path = f"page_{page_num + 1}.png"  # Název souboru obrázku
        page_path = docx_file.replace(".docx", f"_{page_num + 1}.png")  # Cesta k obrázku v dokumentu DOCX
        pixmap = page.get_pixmap()
        pixmap.save(img_path)  # Export stránky PDF jako PNG obrázku
        doc.add_picture(img_path, width=Inches(7))  # Vložení obrázku do dokumentu DOCX

    # Uložení dokumentu DOCX
    doc.save(docx_file)

def convert_all_pdfs_in_folder(folder_path):
    # Procházení souborů ve složce
    for filename in os.listdir(folder_path):
        if filename.lower().endswith(".pdf"):
            pdf_file = os.path.join(folder_path, filename)
            docx_file = os.path.join(folder_path, filename.replace(".pdf", ".docx"))
            
            # Zavolejte funkci pro vytvoření DOCX s obsahem z PDF (lze zvolit jednu z funkcí)
          
            # create_docx_with_pdf_content(pdf_file, docx_file)   # Prevede PDF do DOCX jako text 
            insert_pdf_pages_as_images(pdf_file, docx_file)  # Prevede PDF do DOCX jako obrazek/stranka

# Cesta ke složce s PDF soubory
folder_path = r"C:\Users\ijttr\OneDrive\Dokumenty\OCEŇOVÁNÍ\_IJK\FMP SICAV\2024\LV"

# Funkce pro konverzi všech PDF souborů ve složce
convert_all_pdfs_in_folder(folder_path)
